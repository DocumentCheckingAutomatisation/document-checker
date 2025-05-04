import re
from typing import Dict, Any, List, Optional

from dedoc import DedocManager
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor


class DocxParser:
    def __init__(self, docx_file_path: str):
        self.docx_file_path = docx_file_path
        self.errors = []
        self.parsed_document = self.run_parse()
        self.serialised_document = self.init_dedoc()

    def init_dedoc(self):
        manager = DedocManager()
        result = manager.parse(self.docx_file_path, {"document_type": "diploma"})
        serialised_doc = result.to_api_schema().model_dump()
        return serialised_doc

    def run_parse(self) -> Dict[str, Any]:
        doc = Document(self.docx_file_path)

        return {
            "structure": self.parse_structure(doc),
            "bold_intro_words": self.parse_intro(doc),
            # "text": self.parse_text(doc),
            "lists": self.parse_lists(doc),
            "pictures": self.parse_pictures(doc),
            "tables": self.parse_tables(doc),
            "appendices": self.parse_appendices(),
            "bibliography": self.parse_bibliography(doc)
        }

    def parse_structure(self, doc: Document) -> Dict[str, Any]:
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        numbered_chapters = []
        unnumbered_chapters = []
        sections = []
        common_text = []
        current_chapter = None

        for para in paragraphs:
            text = para.text.strip()

            # Нумерованные заголовки
            if re.match(r"^\d+(\.\d+)?", text):
                if re.match(r"^\d+\s", text):  # Глава, например "1 Введение"
                    parts = text.split(maxsplit=1)
                    chapter_number = parts[0]
                    chapter_title = parts[1] if len(parts) > 1 else ""
                    formatted_title = f"{chapter_number} глава '{chapter_title}'"

                    para_info = self.extract_paragraph_info(para)
                    para_info["formatted_title"] = formatted_title
                    numbered_chapters.append(para_info)

                    current_chapter = chapter_number

                elif re.match(r"^\d+\.\d+", text):  # Раздел, например "1.1 Название"
                    section_info = self.extract_paragraph_info(para)
                    section_info["chapter_number"] = current_chapter
                    sections.append(section_info)

            else:  # Ненумерованные
                lowered = text.lower()
                if "выводы по главе" in lowered:
                    section_info = self.extract_paragraph_info(para)
                    section_info["chapter_number"] = current_chapter
                    sections.append(section_info)

                elif "технико-экономическое обоснование" in lowered:
                    section_info = self.extract_paragraph_info(para)
                    section_info["chapter_number"] = current_chapter
                    sections.append(section_info)


                elif (text.isupper()
                      and not re.search(r"[.:;!?–—-]$", text)
                      and not re.search(r"[–—-]", text)):
                    unnumbered_chapters.append(self.extract_paragraph_info(para))
                else:
                    # Простой абзац
                    common_text.append(self.extract_paragraph_info(para))

        return {
            "numbered_chapters": numbered_chapters,
            "unnumbered_chapters": unnumbered_chapters,
            "sections": sections,
            "common_text": common_text
        }

    def extract_paragraph_info(self, para) -> Dict[str, Any]:
        run = para.runs[0] if para.runs else None
        return {
            "content": para.text.strip(),
            "font_size": run.font.size.pt if run and run.font.size else None,
            "font_name": run.font.name if run else None,
            "font_color": self.get_font_color(run),
            "line_spacing": para.paragraph_format.line_spacing,
            "alignment": para.alignment if para.alignment else None,
            "left_indent": para.paragraph_format.left_indent.pt if para.paragraph_format.left_indent else None,
            "bold": run.bold if run else None,
            "uppercase": para.text.isupper(),
            "italic": run.italic if run else None,
            "underline": run.underline if run else None
        }

    def get_font_color(self, run) -> Optional[str]:
        if run and run.font.color and run.font.color.rgb:
            rgb: RGBColor = run.font.color.rgb
            return str(rgb)
        return None

    def parse_intro(self, doc: Document) -> List[str]:
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        inside_intro = False
        bold_words = []

        for para in paragraphs:
            text = para.text.strip().lower()
            if "введение" in text:
                inside_intro = True
                continue
            if inside_intro:
                if para.style.name.lower().startswith('heading'):
                    break  # Введение закончилось
                for run in para.runs:
                    if run.bold:
                        bold_words.append(run.text.strip().lower())

        return bold_words

    def parse_lists(self, doc: Document) -> List[Dict[str, Any]]:
        lists = []
        current_list = None
        intro_candidate = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue


            numbered_match = re.match(r"^(\d+)[\.\)]\s+(.+)", text)

            bulleted_match = re.match(r"^[•\-–—]\s+(.+)", text)

            if numbered_match:
                item = {
                    "type": "numbered",
                    "number": int(numbered_match.group(1)),
                    "text": numbered_match.group(2),
                    "full": text
                }

                if current_list is None:
                    current_list = {
                        "type": "numbered",
                        "intro": intro_candidate,
                        "items": [item]
                    }
                else:
                    current_list["items"].append(item)
                continue

            elif bulleted_match:
                item = {
                    "type": "bulleted",
                    "text": bulleted_match.group(1),
                    "full": text
                }

                if current_list is None:
                    current_list = {
                        "type": "bulleted",
                        "intro": intro_candidate,
                        "items": [item]
                    }
                else:
                    current_list["items"].append(item)
                continue

            else:
                if current_list:
                    lists.append(current_list)
                    current_list = None

                if text.endswith(":") or text.endswith("."):
                    intro_candidate = text
                else:
                    intro_candidate = None

        if current_list:
            lists.append(current_list)

        return lists

    def parse_pictures(self, doc: Document) -> Dict[str, List[Dict[str, str]]]:
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        picture_references = []
        picture_captions = []

        # Пары "паттерн — флаг захвата номера рисунка"
        reference_patterns = [
            (r"\(рис\.?\s*(\d+)\)", re.IGNORECASE),  # (Рис. 4)
            (r"\(см\.?\s*рис\.?\s*(\d+)\)", re.IGNORECASE),  # (см. рис. 4)
            (r"на\s+рисунке\s+(\d+)", re.IGNORECASE),  # на рисунке 4
        ]

        caption_pattern = r"^Рисунок\s+(\d+)\s*([-–—])\s*(.+)"

        for i, text in enumerate(paragraphs):
            # Поиск ссылок
            for pattern, flags in reference_patterns:
                for match in re.finditer(pattern, text, flags=flags):
                    picture_references.append({
                        "ref_text": match.group(0),
                        "ref_number": match.group(1),
                        "paragraph": text
                    })

            match = re.match(caption_pattern, text)
            if match:
                picture_captions.append({
                    "figure_number": match.group(1),
                    "dash": match.group(2),
                    "caption": match.group(3),
                    "full_text": text
                })

        return {
            "references": picture_references,
            "captions": picture_captions
        }

    def parse_tables(self, doc: Document) -> Dict[str, List[Dict[str, Any]]]:
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        table_references = []
        table_captions = []
        tables = []

        reference_patterns = [
            (r"\(табл\.?\s*(\d+)\)", re.IGNORECASE),  # (табл. 1)
            (r"\(см\.?\s*табл\.?\s*(\d+)\)", re.IGNORECASE),  # (см. табл. 1)
            (r"в\s+таблице\s+(\d+)", re.IGNORECASE),  # в таблице 1
            (r"из\s+таблицы\s+(\d+)", re.IGNORECASE),  # из таблицы 1
        ]

        for para in paragraphs:
            text = para.text.strip()
            for pattern, flags in reference_patterns:
                for match in re.finditer(pattern, text, flags=flags):
                    table_references.append({
                        "ref_text": match.group(0),
                        "ref_number": match.group(1),
                        "paragraph": text
                    })

        for i in range(len(paragraphs) - 1):
            current = paragraphs[i]
            next_para = paragraphs[i + 1]
            if re.match(r"Таблица\s+\d+", current.text.strip(), re.IGNORECASE):
                if current.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT and next_para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    table_captions.append({
                        "table_number": re.findall(r"\d+", current.text)[0],
                        "title": next_para.text.strip(),
                        "raw_text": f"{current.text.strip()} / {next_para.text.strip()}"
                    })

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append({"table": table_data})

        return {
            "references": table_references,
            "captions": table_captions,
            "tables": tables
        }

    def parse_bibliography(self, doc: Document) -> Dict[str, Any]:
        references = set()
        bibliography_items = []
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        in_bibliography = False
        counter = 1

        for para in paragraphs:
            # Сбор ссылок вида [1], [2]
            refs = re.findall(r"\[\d+\]", para)
            references.update(refs)

            lowered = para.lower()

            # Запуск сбора библиографии
            if not in_bibliography and "список использованных источников" in lowered:
                in_bibliography = True
                continue

            # Сбор элементов библиографии
            if in_bibliography:
                if para == para.upper():  # Новый заголовок = выход
                    in_bibliography = False
                    continue

                number = para[0] if para and para[0].isdigit() else None
                bibliography_items.append({
                    "number": counter,
                    "content": para
                })
                counter += 1

        return {
            "references_in_text": list(references),
            "bibliography": bibliography_items
        }

    def parse_appendices(self) -> Dict[str, List[Dict[str, str]]]:
        doc = Document(self.docx_file_path)
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        appendix_references = []
        appendix_titles = []

        # Ссылки на приложения в тексте
        reference_patterns = [
            (r"\(прил\.?\s*([А-ЯA-Z])\)", re.IGNORECASE),  # (прил. А)
            (r"\(см\.?\s*прил\.?\s*([А-ЯA-Z])\)", re.IGNORECASE),  # (см. прил. А)
            (r"в\s+приложении\s+([А-ЯA-Z])", re.IGNORECASE),  # в приложении А
            (r"из\s+приложения\s+([А-ЯA-Z])", re.IGNORECASE),  # из приложения А
        ]

        for para in paragraphs:
            text = para.text.strip()
            for pattern, flags in reference_patterns:
                for match in re.finditer(pattern, text, flags=flags):
                    appendix_references.append({
                        "ref_text": match.group(0),
                        "ref_letter": match.group(1).upper(),
                        "paragraph": text
                    })

        # Названия приложений
        for i in range(len(paragraphs) - 1):
            current = paragraphs[i]
            next_para = paragraphs[i + 1]
            match = re.match(r"Приложение\s+([А-ЯA-Z])", current.text.strip(), re.IGNORECASE)

            if match:
                if current.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT and next_para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    appendix_titles.append({
                        "appendix_letter": match.group(1).upper(),
                        "title": next_para.text.strip(),
                        "raw_text": f"{current.text.strip()} / {next_para.text.strip()}"
                    })

        return {
            "references": appendix_references,
            "titles": appendix_titles
        }

# import re
# from typing import Dict, Any, List
#
# from dedoc import DedocManager
# from docx import Document
#
#
# class DocxParser:
#     def __init__(self, docx_file_path):
#         self.docx_file_path = docx_file_path
#         self.errors = []
#         self.parsed_document = self.run_parse()
#
#     def run_parse(self) -> Dict[str, Any]:
#         # manager = DedocManager()
#         # result = manager.parse(self.docx_file_path, {"document_type": "diploma"})
#         # serialised_doc = result.to_api_schema().model_dump()
#         doc = Document(self.docx_file_path)
#
#         return {"structure": self.parse_structure(doc),
#                 "bold_intro_words": self.parse_intro(doc),
#                 "text": self.parse_text(doc),
#                 "lists": self.parse_lists(doc),
#                 "tables": self.parse_tables(doc),
#                 "pictures": self.parse_pictures(doc),
#                 "literature_links": self.parse_lit_links(doc)
#                 }
#
#     def parse_structure(self, doc: Document) -> Dict[str, Any]:
#         # парсинг нумерованных (1) и ненумерованных глав (содержание введение и тд), нумерованных (1.1) и ненумерованных разделов (выводы по главе)
#         # у разделов должно быть указано в какой они главе находятся
#         paragraphs = [p for p in doc.paragraphs if p.text.strip()]
#
#         return {
#             "numbered_chapters": {"content": ..., "font_size": ..., "font_name": ..., "font_color": ...,
#                                   "line_spacing": ..., "alignment": ..., "left_indent": ..., "bold": ...,
#                                   "uppercase": ..., "italic": ..., "underline": ...}, # добавляй еще если есть
#             "unnumbered_chapters": ...,
#             "numbered_sections": ...,
#             "unnumbered_sections": ...
#         }
#
#     def parse_intro(self, doc: Document) -> List[str]:
#         paragraphs = [p for p in doc.paragraphs if p.text.strip()]
#         inside_intro = False
#         bold_words = []
#
#         for para in paragraphs:
#             text = para.text.strip().lower()
#             if "введение" in text:
#                 inside_intro = True
#                 continue
#             if inside_intro:
#                 if para.style.name.lower().startswith('heading'):
#                     break  # Введение закончилось
#
#                 for run in para.runs:
#                     if run.bold:
#                         bold_words.append(run.text.lower())
#
#         return bold_words
#
#     def parse_text(self, doc: Document) -> Dict[str, Any]:
#         # парсинг кусков текста за исключением заголовков глав и разделов
#         return {
#             "content": ...,
#             "font_size": ...,
#             "font_name": ...,
#             "font_color": ...,
#             "line_spacing": ...,
#             "alignment": ...,
#             "left_indent": ...,
#             "bold": ...,
#             "uppercase": [run.text.isupper() for p in doc.paragraphs for run in p.runs],
#             "italic": ...,
#             "underline": ...
#         }
#
#     def parse_lists(self, serialised_doc):
#         # парсинг списков: вводное предложение (то есть 1 предложение до элементов списка) и сами элементы списка
#         pass
#
#     def parse_tables(self, serialised_doc):
#         pass
#
#     def parse_pictures(self, serialised_doc):
#         # парсинг рисунков их подписей и ссылок на них и свойств (отступы размеры шрифта и тд)
#         pass
#
#     def parse_lit_links(self, serialised_doc):
#         # нахождение ссылки в формате [1] на элемент в списке использованных источников и самого элемента
#         pass


#
# from typing import Dict, Any, List
#
# from dedoc import DedocManager
#
#
# class DocxParser:
#     def __init__(self, docx_file_path):
#         self.docx_file_path = docx_file_path
#         self.errors = []
#         self.parsed_document = self.run_parse()
#
#     def run_parse(self) -> Dict[str, Any]:
#         manager = DedocManager()
#         result = manager.parse(self.docx_file_path, {"document_type": "diploma"})
#         serialised_doc = result.to_api_schema().model_dump()
#
#         return {"structure": self.parse_structure(serialised_doc),
#                 "fonts": self.parse_font_details(serialised_doc)
#                 }
#
#     def parse_structure(self, serialised_doc: Dict[str, Any]) -> Dict[str, Any]:
#         chapters = serialised_doc["content"]["structure"]
#         return {"found_chapters": chapters}
#
#     def parse_font_details(self, serialised_doc: Dict[str, Any]) -> Dict[str, Any]:
#         fonts = {}
#         for annotation in serialised_doc["content"]["structure"]["subparagraphs"]:
#             title = annotation["text"]
#             font_size, font_bold = self.extract_font_details(annotation.get("annotations", []))
#             fonts[title] = {"size": font_size, "bold": font_bold}
#         return fonts
#
#     @staticmethod
#     def extract_font_details(annotations: List[Dict[str, Any]]):
#         font_size = None
#         font_bold = False
#         for annotation in annotations:
#             if annotation["name"] == "size":
#                 font_size = annotation["value"]
#             elif annotation["name"] == "bold":
#                 font_bold = True
#         return font_size, font_bold
